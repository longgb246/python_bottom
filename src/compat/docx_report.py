# -*- coding:utf-8 -*-
# @Author  : 'longguangbin'
# @Contact : lgb453476610@163.com
# @Date    : 2018/12/27
"""  
Usage Of 'docx_report' : 
"""

from __future__ import print_function
import os

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt


def hex_to_rgb(hex_str):
    """hex_to_rgb

    :param hex_str: '#7BF5BE'
    :return: rgb list - [123, 245, 190]
    """
    hex_str = hex_str.replace('#', '')
    rgb_v = map(lambda x: int(x, 16), [hex_str[:2], hex_str[2:4], hex_str[4:6]])
    return rgb_v


def get_rgb_color(color_v):
    """Get the rgb color from hex """
    color_map = {'black': '#000000',
                 'red': '#EB3223',
                 'orange': '#F6C143',
                 'yellow': '#FFFD54',
                 'green': '#4EAC5B',
                 'blue': '#2C70BA',
                 'purple': '#68379A'}
    if color_v in color_map.keys():
        color_v = color_map.get(color_v, '#000000')
    if color_v.startswith('#'):
        color_res = hex_to_rgb(color_v)
    else:
        raise Exception
    return color_res


def set_font_styles(run_ins, styles):
    """设置 word 字体风格

    :param run_ins: Run 对象
    :param dict styles: {'type': u'宋体',
                         'bold': False,
                         'italic': False,
                         'size': 14,
                         'underline': False,
                         'color': 'black',  # '#000000'
                         }
    :return: Run 对象
    """
    if 'type' in styles.keys():
        f_type = styles.get('type')
        run_ins.font.name = f_type
        run_ins._element.rPr.rFonts.set(qn('w:eastAsia'), f_type)
    if 'bold' in styles.keys():
        run_ins.font.bold = styles.get('bold')
    if 'italic' in styles.keys():
        run_ins.font.italic = styles.get('italic')
    if 'underline' in styles.keys():
        run_ins.font.underline = styles.get('underline')
    if 'size' in styles.keys():
        run_ins.font.size = Pt(styles.get('size', 14))
    if 'color' in styles.keys():
        rgb_color = get_rgb_color(styles.get('color'))
        run_ins.color.rgb = RGBColor(*rgb_color)
    return run_ins


class ReportBase(object):
    def __init__(self):
        self.document = Document()
        # style
        self._set_default_styles()

    def _set_default_styles(self):
        """使用默认的风格 """
        self._txt_type = u'宋体'
        self._head_type = u'微软雅黑'
        cn_list = ['Normal', 'TOC Heading', 'Heading 1', 'Title', 'Subtitle', 'Body Text', 'Body Text 2']
        for each_cn in cn_list:
            set_font_styles(self.document.styles[each_cn], styles={'type': self._txt_type})

    def heading(self, content, level):
        """添加文本标题

        :param (unicode or str or list) content: list format - [[content1, style1], [content2, style2], ...]
        :param level:
        :return: Run 对象 or list
        """
        run_ins = self.document.add_heading('', level).add_run(content)
        set_font_styles(run_ins, styles={'type': self._head_type})
        return run_ins

    def text(self, content):
        """添加段落文本，可以输入 string 类型为直接转化为文本，也可以输入 list 分段转化为文本

        :param (unicode or str or list) content: list format - [[content1, style1], [content2, style2], ...]
        :return: Run 对象 or list
        """
        if isinstance(content, (str, unicode)):
            run_ins = self.document.add_paragraph('').add_run(content)
        elif isinstance(content, list):
            run_ins = ''
        else:
            raise Exception()
        return run_ins

    def picture(self, pic_file):
        """添加图片 """
        # pic_file = self.file_path + os.sep + 'caigou.png'
        # height=Inches(3.84), width=Inches(5.76)
        # height=Inches(2.88), width=Inches(5.76)
        self.document.add_picture(pic_file, height=Inches(3.84), width=Inches(5.76))
        pass

    def table(self, data):
        """添加表格

        :param (list or pd.Dataframe) data: list - [[header], [value1], [value2], ...]
        :return:
        """
        if hasattr(data, 'values'):
            columns = data.columns
            data = [columns] + data.values.tolist()
        header = data[0]
        contents = data[1:]
        col_len = len(header)
        table = self.document.add_table(1, col_len)
        table.style = 'Light List Accent 1'
        heading_cells = table.rows[0].cells

        # add heading
        for i in range(col_len):
            heading_cells[i].text = unicode(header[i])

        # add content
        for content in contents:
            cells = table.add_row().cells
            for i, item in enumerate(content):
                cells[i].text = unicode(item)

    def save(self, save_path):
        """保存docx文件 """
        self.document.save(save_path)


import pandas as pd


class EDAReport(ReportBase):
    def __init__(self, path):
        super(EDAReport, self).__init__()
        self.save_path = path
        self.report_name = 'report_word'

    def header_statistic(self):
        self.heading(u'A 品类概况a', 0)
        self.heading(u'a) 数据准备', 1)
        self.text(u"\tSKU总数为{0}个，在分析首日之前上柜的SKU总数为{1}个：".format(122, 222))
        self.text(u"\t1）首日销量预测为空的sku数量为：{0}".format(200))
        a_pd = pd.DataFrame([[1, 2, 3], [4, 5, 6]], columns=['cr', 'mv', 'pe'])
        self.table(a_pd)

    def save_report(self):
        self.save(self.save_path + os.sep + '{0}.docx'.format(self.report_name))

    def run(self):
        self.header_statistic()
        self.save_report()


if __name__ == '__main__':
    # path = sys.argv[1]
    path = '/Users/longguangbin/Work'
    eda_report = EDAReport(path=path)
    eda_report.run()
